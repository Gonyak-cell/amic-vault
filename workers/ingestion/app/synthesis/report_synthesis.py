from dataclasses import dataclass
from io import BytesIO

from docx import Document


@dataclass(frozen=True)
class DdIssue:
    code: str
    title: str
    severity: str
    status: str
    citation_refs: tuple[str, ...]


@dataclass(frozen=True)
class DdRisk:
    code: str
    category: str
    severity: str
    likelihood: str
    mitigation_summary: str | None
    citation_refs: tuple[str, ...]


@dataclass(frozen=True)
class NegotiationIssue:
    issue_id: str
    clause_ref: str
    playbook_rule: str
    redline_summary: str
    status: str
    citation_refs: tuple[str, ...]


def build_dd_report_docx(
    *,
    matter_label: str,
    issues: tuple[DdIssue, ...],
    risks: tuple[DdRisk, ...],
) -> bytes:
    doc = Document()
    doc.add_heading(f"{_safe_text(matter_label, 120)} DD 보고서 초안", 0)
    doc.add_heading("이슈 요약", level=1)
    issue_table = doc.add_table(rows=1, cols=5)
    _set_cells(issue_table.rows[0].cells, ("코드", "제목", "심각도", "상태", "인용"))
    for issue in issues:
        _set_cells(
            issue_table.add_row().cells,
            (
                _safe_text(issue.code, 64),
                _safe_text(issue.title, 240),
                _safe_text(issue.severity, 32),
                _safe_text(issue.status, 32),
                _citation_text(issue.citation_refs),
            ),
        )

    doc.add_heading("리스크 등급표", level=1)
    risk_table = doc.add_table(rows=1, cols=6)
    _set_cells(
        risk_table.rows[0].cells,
        ("코드", "분류", "심각도", "가능성", "완화 요약", "인용"),
    )
    for risk in risks:
        _set_cells(
            risk_table.add_row().cells,
            (
                _safe_text(risk.code, 64),
                _safe_text(risk.category, 64),
                _safe_text(risk.severity, 32),
                _safe_text(risk.likelihood, 32),
                _safe_text(risk.mitigation_summary or "검토 필요", 500),
                _citation_text(risk.citation_refs),
            ),
        )

    doc.add_heading("인용 목록", level=1)
    for ref in _unique_refs(
        tuple(ref for issue in issues for ref in issue.citation_refs)
        + tuple(ref for risk in risks for ref in risk.citation_refs)
    ):
        doc.add_paragraph(_safe_ref(ref), style="List Bullet")
    return _save_docx(doc)


def build_negotiation_issues_docx(
    *,
    matter_label: str,
    issues: tuple[NegotiationIssue, ...],
) -> bytes:
    doc = Document()
    doc.add_heading(f"{_safe_text(matter_label, 120)} 협상쟁점표", 0)
    table = doc.add_table(rows=1, cols=6)
    _set_cells(table.rows[0].cells, ("조항", "룰", "레드라인", "상태", "식별자", "인용"))
    for issue in issues:
        _set_cells(
            table.add_row().cells,
            (
                _safe_text(issue.clause_ref, 160),
                _safe_text(issue.playbook_rule, 160),
                _safe_text(issue.redline_summary, 500),
                _safe_text(issue.status, 32),
                _safe_text(issue.issue_id, 80),
                _citation_text(issue.citation_refs),
            ),
        )
    doc.add_heading("인용 목록", level=1)
    for ref in _unique_refs(tuple(ref for issue in issues for ref in issue.citation_refs)):
        doc.add_paragraph(_safe_ref(ref), style="List Bullet")
    return _save_docx(doc)


def _set_cells(cells, values: tuple[str, ...]) -> None:
    for cell, value in zip(cells, values, strict=True):
        cell.text = value


def _safe_text(value: str, limit: int) -> str:
    normalized = " ".join(value.split()).strip()
    lowered = normalized.lower()
    if any(token in lowered for token in ("password", "secret", "token")):
        raise ValueError("unsafe synthesis value")
    return normalized[:limit] or "검토 필요"


def _safe_ref(value: str) -> str:
    normalized = _safe_text(value, 160)
    lowered = normalized.lower()
    if any(token in lowered for token in ("body", "content", "snippet", "raw")):
        raise ValueError("unsafe citation ref")
    return normalized


def _citation_text(values: tuple[str, ...]) -> str:
    refs = _unique_refs(values)
    if not refs:
        raise ValueError("at least one citation ref is required")
    return ", ".join(_safe_ref(ref) for ref in refs)


def _unique_refs(values: tuple[str, ...]) -> tuple[str, ...]:
    seen: set[str] = set()
    output: list[str] = []
    for value in values:
        ref = value.strip()
        if ref and ref not in seen:
            seen.add(ref)
            output.append(ref)
    return tuple(output)


def _save_docx(doc: Document) -> bytes:
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
