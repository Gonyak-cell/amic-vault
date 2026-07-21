from io import BytesIO

import pytest
from docx import Document

from app.synthesis.report_synthesis import (
    DdIssue,
    DdRisk,
    NegotiationIssue,
    build_dd_report_docx,
    build_negotiation_issues_docx,
)


def test_build_dd_report_docx_contains_required_sections_and_citations():
    body = build_dd_report_docx(
        matter_label="Project Atlas",
        issues=(
            DdIssue(
                code="DD.ISSUE.1",
                title="Change of control consent is missing",
                severity="high",
                status="open",
                citation_refs=("dd_issue:11111111-1111-4111-8111-111111111111",),
            ),
        ),
        risks=(
            DdRisk(
                code="DD.RISK.1",
                category="legal",
                severity="critical",
                likelihood="medium",
                mitigation_summary="Obtain consent before closing",
                citation_refs=("dd_risk:22222222-2222-4222-8222-222222222222",),
            ),
        ),
    )

    doc = Document(BytesIO(body))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    tables = [[cell.text for row in table.rows for cell in row.cells] for table in doc.tables]
    flattened = "\n".join(cell for table in tables for cell in table)
    assert "Project Atlas DD 보고서 초안" in text
    assert "이슈 요약" in text
    assert "리스크 등급표" in text
    assert "인용 목록" in text
    assert "Change of control consent is missing" in flattened
    assert "Obtain consent before closing" in flattened
    assert "dd_issue:11111111-1111-4111-8111-111111111111" in flattened
    assert "dd_risk:22222222-2222-4222-8222-222222222222" in text


def test_build_negotiation_issues_docx_contains_issue_table():
    body = build_negotiation_issues_docx(
        matter_label="Project Atlas",
        issues=(
            NegotiationIssue(
                issue_id="NI-1",
                clause_ref="Section 8.2",
                playbook_rule="cap-liability",
                redline_summary="Counterparty removed indirect damages exclusion",
                status="open",
                citation_refs=("negotiation_issue:33333333-3333-4333-8333-333333333333",),
            ),
        ),
    )

    doc = Document(BytesIO(body))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    flattened = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "Project Atlas 협상쟁점표" in text
    assert "Section 8.2" in flattened
    assert "cap-liability" in flattened
    assert "Counterparty removed indirect damages exclusion" in flattened
    assert "negotiation_issue:33333333-3333-4333-8333-333333333333" in text


def test_report_synthesis_rejects_uncited_or_unsafe_rows():
    with pytest.raises(ValueError):
        build_dd_report_docx(
            matter_label="Project Atlas",
            issues=(
                DdIssue(
                    code="DD.ISSUE.2",
                    title="Missing citation",
                    severity="medium",
                    status="open",
                    citation_refs=(),
                ),
            ),
            risks=(),
        )
    with pytest.raises(ValueError):
        build_negotiation_issues_docx(
            matter_label="Project Atlas",
            issues=(
                NegotiationIssue(
                    issue_id="NI-2",
                    clause_ref="Section 9",
                    playbook_rule="unsafe",
                    redline_summary="Contains password placeholder",
                    status="open",
                    citation_refs=("negotiation_issue:44444444-4444-4444-8444-444444444444",),
                ),
            ),
        )
